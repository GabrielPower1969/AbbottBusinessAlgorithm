from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
import os
import re
from MyLogger import logger

class WordWriter:
    def __init__(self, template_file, data_dict):
        self.template_file = template_file
        self.data_dict = data_dict

    def parse_style(self, style_string):
        """
        Parse style information from a string.
        :param style_string: The style string, e.g., "font=Arial, size=24, bold=True, color=FF0000".
            Red: FF0000
            Green: 00FF00
            Blue: 0000FF
            Black: 000000
            White: FFFFFF
            Yellow: FFFF00
            Cyan: 00FFFF
            Magenta: FF00FF
            Gray: 808080
            Orange: FFA500
            Purple: 800080
            Brown: A52A2A
            You can use these values in your style strings to set the color
            of text in your Word document.
            For example, to set the text color to red, you would use color=FF0000.


        :return: A dictionary containing style information.

        """
        style_info = {}
        styles = style_string.split(", ")
        for style in styles:
            key, value = style.split("=")
            if key == 'size':
                value = int(value)
            elif key == 'bold':
                value = value == 'True'
            elif key == 'color':
                value = RGBColor.from_string(value)
            style_info[key] = value
        return style_info

    def apply_style(self, run, style_info):
        """
        Apply style to a run object.
        :param run: The run object to apply styles to.
        :param style_info: A dictionary containing style information.
        """
        if 'font' in style_info:
            run.font.name = style_info['font']
        if 'size' in style_info:
            run.font.size = Pt(style_info['size'])
        if 'bold' in style_info:
            run.bold = style_info['bold']
        if 'color' in style_info:
            run.font.color.rgb = style_info['color']

    def write_to_word(self, output_file):
        """
        Write the order data to a Word document with the specified template.
        The document is set to A5 size, landscape orientation, with narrow margins.
        The template uses placeholders to mark where data should be inserted.
        """
        if self.template_file == "":
            print("No Abbott product template is available for this order")
            return

        # Read the template file content
        with open(self.template_file, 'r', encoding = 'utf-8') as infile:
            print(f"Reading template file: {self.template_file}")
            lines = infile.read().splitlines()

        document = Document()

        # Set the document to A5 size, landscape, and narrow margins
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(5.83)
        section.orientation = WD_ORIENT.LANDSCAPE

        # Set narrow margins
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # Add template content to the Word document
        for line in lines:
            p = document.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Set line spacing to 1.0
            parts = re.split(r'(\{.*?\})', line)
            for part in parts:
                if part.startswith('{') and part.endswith('}'):
                    content = part[1:-1] # get the content between the curly braces
                    if ':' in content:
                        key, style_string = content.split(': ')
                        style_info = self.parse_style(style_string)
                        if key.startswith('*content*'):
                            fixed_text = key[len('*content*'):].strip()
                            run = p.add_run(fixed_text)
                            self.apply_style(run, style_info)
                        else:
                            if key in self.data_dict:
                                value = str(self.data_dict[key])
                                run = p.add_run(value)
                                self.apply_style(run, style_info)
                            else:
                                logger.warning(f"The template placeholder is incorrectly configured‘{key}’inexistence")
                                break
                    else:
                        run = p.add_run(part)
                else:
                    run = p.add_run(part)

        # Save the document to the specified output file
        document.save(output_file)
        print(f"Order form saved as {output_file}")

    def merge_word_documents(self, files, output_file):
        """
        Merge multiple Word documents into a single document

        Args:
            files (list): List of file paths to merge
            output_file (str): Output file path for merged document
        """
        # 直接使用第一个非空的文档作为基础
        merged_doc = None
        first_doc_index = -1

        # 找到第一个有内容的文档
        for i, file in enumerate(files):
            doc = Document(file)
            if any(paragraph.text.strip() for paragraph in doc.paragraphs):
                merged_doc = doc
                first_doc_index = i
                break

        if merged_doc is None:
            logger.error("No valid document found to merge")
            return

        # 从下一个文档开始合并
        for file in files[first_doc_index + 1:]:
            sub_doc = Document(file)
            # 直接添加内容，不添加分页符
            if any(paragraph.text.strip() for paragraph in sub_doc.paragraphs):
                for element in sub_doc.element.body:
                    merged_doc.element.body.append(element)

        # 保存合并后的文档
        merged_doc.save(output_file)

        # 清理临时文件
        for file in files[1:]:
            try:
                os.remove(file)
            except Exception as e:
                logger.warning(f"Failed to remove temporary file {file}: {str(e)}")